"""
Production-grade resume parser with deterministic extraction.
NO HALLUCINATIONS - All data must come from the resume text.
"""
import re
import os
from typing import Optional, List, Dict, Any, Tuple
from datetime import datetime
from dateutil import parser as date_parser
import dateparser
import pdfplumber
import fitz  # PyMuPDF
from docx import Document as DocxDocument


# ═══════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(file_path: str) -> Tuple[str, float]:
    """
    Extract text from PDF using multiple methods for robustness.
    Returns: (text, confidence_score)
    """
    text = ""
    confidence = 0.0
    
    try:
        # Method 1: pdfplumber (best for structured PDFs)
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if len(text.strip()) > 100:
            confidence = 0.95
            return text.strip(), confidence
    except Exception as e:
        print(f"pdfplumber extraction failed: {e}")
    
    try:
        # Method 2: PyMuPDF (fallback for complex PDFs)
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        
        if len(text.strip()) > 100:
            confidence = 0.85
            return text.strip(), confidence
    except Exception as e:
        print(f"PyMuPDF extraction failed: {e}")
    
    # Low confidence if extraction yielded minimal text
    if len(text.strip()) < 50:
        confidence = 0.3
    
    return text.strip(), confidence


def extract_text_from_docx(file_path: str) -> Tuple[str, float]:
    """Extract text from DOCX with confidence scoring."""
    text = ""
    try:
        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        confidence = 0.95 if len(text.strip()) > 100 else 0.5
        return text.strip(), confidence
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return "", 0.0


def extract_text(file_path: str, file_type: str) -> Tuple[str, float]:
    """
    Route to correct extractor.
    Returns: (text, confidence_score)
    """
    if file_type in ["pdf", "application/pdf"]:
        return extract_text_from_pdf(file_path)
    elif file_type in ["docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        return extract_text_from_docx(file_path)
    elif file_type in ["txt", "text/plain"]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return text, 0.9
        except:
            return "", 0.0
    return "", 0.0


# ═══════════════════════════════════════════════════════════════════════════
# PATTERN-BASED EXTRACTORS
# ═══════════════════════════════════════════════════════════════════════════

EMAIL_RE = re.compile(r"\b[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+\b")
PHONE_RE = re.compile(r"(\+?\d[\d\s\-().]{7,15}\d)")

# Comprehensive tech skills database
TECH_SKILLS = {
    "languages": ["python", "java", "javascript", "typescript", "go", "golang", "rust", "c++", "cpp", "c#", "csharp", "ruby", "php", "scala", "kotlin", "swift", "r", "perl", "dart", "elixir", "clojure", "haskell"],
    "frameworks": ["django", "fastapi", "flask", "spring", "springboot", "react", "reactjs", "angular", "vue", "vuejs", "node.js", "nodejs", "express", "expressjs", "laravel", "rails", "nextjs", "next.js", "nestjs", "svelte", "ember", "backbone"],
    "databases": ["postgresql", "postgres", "mysql", "mongodb", "mongo", "redis", "elasticsearch", "cassandra", "dynamodb", "sqlite", "oracle", "mssql", "sql server", "neo4j", "clickhouse", "mariadb", "couchdb", "influxdb"],
    "cloud": ["aws", "amazon web services", "gcp", "google cloud", "azure", "microsoft azure", "ec2", "s3", "lambda", "rds", "cloudfront", "gke", "aks", "eks", "cloud functions", "app engine"],
    "devops": ["docker", "kubernetes", "k8s", "jenkins", "github actions", "gitlab ci", "circleci", "terraform", "ansible", "helm", "argocd", "puppet", "chef", "vagrant", "ci/cd"],
    "ml": ["tensorflow", "pytorch", "scikit-learn", "sklearn", "pandas", "numpy", "huggingface", "langchain", "openai", "mlflow", "airflow", "keras", "xgboost", "lightgbm", "spacy", "nltk"],
    "tools": ["git", "github", "gitlab", "linux", "kafka", "rabbitmq", "graphql", "rest api", "grpc", "celery", "nginx", "apache", "prometheus", "grafana", "datadog", "splunk", "jira", "confluence"],
    "frontend": ["html", "css", "sass", "scss", "tailwind", "bootstrap", "material-ui", "mui", "webpack", "vite", "babel", "redux", "mobx", "jquery"],
    "testing": ["jest", "pytest", "junit", "selenium", "cypress", "mocha", "chai", "jasmine", "testng", "cucumber"],
    "mobile": ["react native", "flutter", "android", "ios", "swift", "kotlin", "xamarin"],
}

ALL_SKILLS = []
for category in TECH_SKILLS.values():
    ALL_SKILLS.extend(category)

EDUCATION_KEYWORDS = {
    "phd": ["phd", "ph.d", "ph. d", "doctorate", "doctor of philosophy", "doctoral"],
    "masters": ["m.tech", "m.s.", "m.s", "msc", "m.sc", "m.e.", "mba", "master of", "masters", "m.a.", "mca"],
    "bachelors": ["b.tech", "b.e.", "b.sc", "b.s.", "bachelor of", "bachelors", "undergraduate", "b.a.", "bca", "bba"],
    "diploma": ["diploma", "polytechnic", "associate degree"],
}

PREMIUM_INSTITUTES = [
    "iit", "iim", "nit", "bits pilani", "bits", "iiit", "iisc", "iiser",
    "mit", "stanford", "oxford", "cambridge", "harvard", "cmu", 
    "carnegie mellon", "berkeley", "caltech", "princeton", "yale", 
    "cornell", "columbia", "upenn", "duke", "michigan", "georgia tech"
]


def extract_email(text: str) -> Optional[str]:
    """Extract email with validation."""
    matches = EMAIL_RE.findall(text)
    # Filter out common false positives
    valid_emails = [m for m in matches if not any(x in m.lower() for x in ["example.com", "test.com", "sample.com"])]
    return valid_emails[0] if valid_emails else None


def extract_phone(text: str) -> Optional[str]:
    """Extract phone number."""
    matches = PHONE_RE.findall(text)
    if matches:
        # Clean up the phone number
        phone = matches[0].strip()
        # Remove excessive spaces
        phone = re.sub(r'\s+', ' ', phone)
        return phone
    return None


def extract_skills(text: str) -> List[str]:
    """
    Extract technical skills using word boundary matching.
    Returns deduplicated list of found skills.
    """
    text_lower = text.lower()
    found = set()
    
    for skill in ALL_SKILLS:
        # Use word boundary for accurate matching
        pattern = r'\b' + re.escape(skill.replace('.', r'\.')) + r'\b'
        if re.search(pattern, text_lower):
            # Normalize skill name
            normalized = skill.replace('.js', '').replace('js', '').strip()
            found.add(skill)
    
    return sorted(list(found))


def extract_name(text: str) -> Tuple[str, float]:
    """
    Extract candidate name with confidence scoring.
    Returns: (name, confidence)
    """
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    
    if not lines:
        return "Unknown", 0.0
    
    # Try first few lines
    for line in lines[:5]:
        # Skip lines with common header keywords
        skip_keywords = ["resume", "curriculum vitae", "cv", "profile", "contact", "email", "phone", "address", "http", "@"]
        if any(kw in line.lower() for kw in skip_keywords):
            continue
        
        # Name should be 2-5 words, mostly alphabetic
        words = line.split()
        if 2 <= len(words) <= 5:
            # Check if mostly alphabetic
            alpha_ratio = sum(c.isalpha() or c.isspace() for c in line) / max(len(line), 1)
            if alpha_ratio > 0.7:
                return line, 0.85
    
    # Fallback: use first non-empty line
    first = lines[0]
    if len(first.split()) <= 6:
        return first, 0.5
    
    return "Unknown", 0.0


def parse_date(date_str: str) -> Optional[datetime]:
    """
    Parse date string to datetime object.
    Handles various formats: "Jan 2020", "January 2020", "2020", "01/2020", etc.
    """
    if not date_str:
        return None
    
    # Handle "Present", "Current", "Now"
    if date_str.lower() in ["present", "current", "now", "ongoing"]:
        return datetime.now()
    
    try:
        # Try dateparser first (handles most formats)
        parsed = dateparser.parse(date_str, settings={'PREFER_DAY_OF_MONTH': 'first'})
        if parsed:
            return parsed
    except:
        pass
    
    try:
        # Try python-dateutil
        parsed = date_parser.parse(date_str, fuzzy=True)
        return parsed
    except:
        pass
    
    # Try to extract just year
    year_match = re.search(r'\b(19|20)\d{2}\b', date_str)
    if year_match:
        try:
            year = int(year_match.group(0))
            return datetime(year, 1, 1)
        except:
            pass
    
    return None


def calculate_duration_months(start_date: Optional[datetime], end_date: Optional[datetime]) -> int:
    """Calculate duration in months between two dates."""
    if not start_date:
        return 0
    
    if not end_date:
        end_date = datetime.now()
    
    # Calculate month difference
    months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
    return max(months, 0)


def extract_work_experience(text: str) -> Tuple[List[Dict[str, Any]], float]:
    """
    Extract work experience with proper date parsing and duration calculation.
    Returns: (experience_list, confidence_score)
    """
    experience = []
    confidence = 0.0
    
    # Find experience section
    exp_section_pattern = r'(?:EXPERIENCE|WORK EXPERIENCE|EMPLOYMENT|PROFESSIONAL EXPERIENCE|WORK HISTORY)(.*?)(?:EDUCATION|SKILLS|PROJECTS|CERTIFICATIONS|$)'
    exp_match = re.search(exp_section_pattern, text, re.IGNORECASE | re.DOTALL)
    
    if not exp_match:
        # Try to find experience entries without explicit section
        exp_text = text
        confidence = 0.5
    else:
        exp_text = exp_match.group(1)
        confidence = 0.8
    
    # Pattern for date ranges
    date_pattern = re.compile(
        r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|'
        r'\d{1,2}[/-]\d{4}|'
        r'\d{4})'
        r'\s*[-–—to]+\s*'
        r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|'
        r'\d{1,2}[/-]\d{4}|'
        r'\d{4}|'
        r'Present|Current|Now|Ongoing)',
        re.IGNORECASE
    )
    
    # Pattern for company and role
    # Look for lines with company names (often in title case or all caps)
    company_role_pattern = re.compile(
        r'([A-Z][A-Za-z\s&.,\'-]+(?:Inc|Ltd|Pvt|Corp|Corporation|Technologies|Solutions|Systems|Labs|Software|Services|Consulting|Group|Company)?)'
        r'\s*[|·•–-]?\s*'
        r'([A-Z][A-Za-z\s]+(?:Engineer|Developer|Manager|Analyst|Designer|Architect|Lead|Senior|Junior|Intern|Consultant|Specialist|Director|VP|Head|Associate)?)',
        re.MULTILINE
    )
    
    # Find all date ranges
    date_matches = list(date_pattern.finditer(exp_text))
    
    for match in date_matches[:10]:  # Limit to 10 jobs max
        start_str = match.group(1)
        end_str = match.group(2)
        
        start_date = parse_date(start_str)
        end_date = parse_date(end_str)
        is_current = end_str.lower() in ["present", "current", "now", "ongoing"]
        
        if not start_date:
            continue
        
        duration_months = calculate_duration_months(start_date, end_date)
        
        # Try to find company and role near this date
        # Look in surrounding text (100 chars before and after)
        context_start = max(0, match.start() - 200)
        context_end = min(len(exp_text), match.end() + 100)
        context = exp_text[context_start:context_end]
        
        company = None
        role = None
        
        # Try to extract company and role from context
        cr_match = company_role_pattern.search(context)
        if cr_match:
            company = cr_match.group(1).strip()
            role = cr_match.group(2).strip()
        
        experience.append({
            "company": company,
            "role": role,
            "start_date": start_str,
            "end_date": end_str,
            "start_date_parsed": start_date.isoformat() if start_date else None,
            "end_date_parsed": end_date.isoformat() if end_date else None,
            "duration_months": duration_months,
            "is_current": is_current,
        })
    
    # Sort by start date (most recent first)
    experience.sort(key=lambda x: x.get("start_date_parsed", ""), reverse=True)
    
    return experience, confidence


def calculate_total_experience_years(work_experience: List[Dict[str, Any]]) -> float:
    """
    Calculate total years of experience from work history.
    CRITICAL: This is the ONLY way to determine experience - NO ESTIMATION.
    """
    if not work_experience:
        return 0.0
    
    total_months = sum(exp.get("duration_months", 0) for exp in work_experience)
    return round(total_months / 12, 1)


def extract_education(text: str) -> Tuple[Dict[str, Any], float]:
    """
    Extract education details with confidence scoring.
    Returns: (education_dict, confidence)
    
    CRITICAL: Only extract from education section to avoid false matches.
    """
    text_lower = text.lower()
    level = "unknown"
    institution = ""
    confidence = 0.0
    
    # Find education section
    edu_section_pattern = r'(?:EDUCATION|ACADEMIC|QUALIFICATION|EDUCATIONAL BACKGROUND)(.*?)(?:EXPERIENCE|WORK HISTORY|SKILLS|PROJECTS|CERTIFICATIONS|ACHIEVEMENTS|$)'
    edu_match = re.search(edu_section_pattern, text, re.IGNORECASE | re.DOTALL)
    
    # CRITICAL FIX: Only search in education section, not entire resume
    if edu_match:
        search_text = edu_match.group(1)
        search_text_lower = search_text.lower()
        confidence = 0.7  # Found education section
    else:
        # Fallback: search entire text but with lower confidence
        search_text = text
        search_text_lower = text_lower
        confidence = 0.3
    
    # Detect education level
    for lvl, keywords in EDUCATION_KEYWORDS.items():
        for kw in keywords:
            if kw in search_text_lower:
                level = lvl
                confidence = max(confidence, 0.8)
                break
        if level != "unknown":
            break
    
    # CRITICAL FIX: Search for premium institutions ONLY in education section
    for inst in PREMIUM_INSTITUTES:
        # Use word boundary to avoid partial matches
        pattern = r'\b' + re.escape(inst) + r'\b'
        if re.search(pattern, search_text_lower):
            institution = inst.upper()
            confidence = 0.9
            break
    
    if not institution:
        # Try to extract institution name from education section only
        # Pattern 1: "University/College/Institute" with proper name
        inst_pattern = re.compile(
            r'(?:from|at|,\s*|^|\n)([A-Z][A-Za-z\s&.\'-]{2,60}(?:University|College|Institute|School))',
            re.MULTILINE
        )
        inst_match = inst_pattern.search(search_text)
        
        if inst_match:
            institution = inst_match.group(1).strip()
            # Clean up common prefixes
            institution = re.sub(r'^(?:from|at|,)\s*', '', institution, flags=re.IGNORECASE)
            confidence = max(confidence, 0.75)
        else:
            # Pattern 2: Look for acronyms (LPU, VIT, etc.)
            acronym_pattern = re.compile(
                r'\b([A-Z]{2,6})\b(?:\s*[-–]\s*)?(?:University|College|Institute)?',
                re.MULTILINE
            )
            acronym_matches = acronym_pattern.findall(search_text)
            
            # Filter out common false positives
            false_positives = {'USA', 'UK', 'US', 'IT', 'CS', 'BE', 'ME', 'MS', 'MBA', 'PHD', 'GPA', 'CGPA'}
            valid_acronyms = [a for a in acronym_matches if a not in false_positives and len(a) >= 3]
            
            if valid_acronyms:
                institution = valid_acronyms[0]  # Take first valid acronym
                confidence = max(confidence, 0.65)
    
    return {
        "level": level,
        "institution": institution,
    }, confidence


def extract_certifications(text: str) -> List[str]:
    """Extract certifications from resume."""
    cert_keywords = [
        "aws certified", "cka", "ckad", "cks", "google cloud certified", 
        "azure certified", "pmp", "csm", "scrum master", "cissp", 
        "comptia", "ccna", "ccnp", "oracle certified", "microsoft certified",
        "tensorflow certified", "kubernetes certified"
    ]
    
    text_lower = text.lower()
    found_certs = []
    
    for cert in cert_keywords:
        if cert in text_lower:
            found_certs.append(cert.title())
    
    return found_certs


def detect_keyword_stuffing(text: str, skills: List[str]) -> bool:
    """
    Flag resumes with excessive skill keyword density.
    This indicates potential resume manipulation.
    """
    word_count = len(text.split())
    if word_count == 0:
        return False
    
    # Check skill density
    skill_density = len(skills) / max(word_count / 100, 1)
    
    # Check for repeated skills
    text_lower = text.lower()
    repeated_count = sum(1 for skill in skills if text_lower.count(skill.lower()) > 5)
    
    return skill_density > 15 or len(skills) > 80 or repeated_count > 10


def detect_duplicate(text: str, existing_texts: List[str]) -> bool:
    """
    Detect duplicate resumes using email and text similarity.
    """
    if not existing_texts:
        return False
    
    # Primary check: email match
    email = extract_email(text)
    if email:
        for existing in existing_texts:
            if extract_email(existing) == email:
                return True
    
    # Secondary check: high text overlap
    words = set(text.lower().split())
    if len(words) == 0:
        return False
    
    for existing in existing_texts:
        existing_words = set(existing.lower().split())
        if len(existing_words) == 0:
            continue
        
        overlap = len(words & existing_words) / len(words | existing_words)
        if overlap > 0.85:
            return True
    
    return False


def calculate_parsing_confidence(
    extraction_confidence: float,
    name_confidence: float,
    has_email: bool,
    has_phone: bool,
    skills_count: int,
    work_exp_confidence: float,
    edu_confidence: float,
) -> float:
    """
    Calculate overall parsing confidence score.
    This helps identify resumes that need manual review.
    """
    score = 0.0
    
    # Text extraction quality (30%)
    score += extraction_confidence * 0.3
    
    # Name extraction (15%)
    score += name_confidence * 0.15
    
    # Contact info (20%)
    if has_email:
        score += 0.15
    if has_phone:
        score += 0.05
    
    # Skills extraction (15%)
    if skills_count > 0:
        score += min(skills_count / 10, 1.0) * 0.15
    
    # Work experience (10%)
    score += work_exp_confidence * 0.1
    
    # Education (10%)
    score += edu_confidence * 0.1
    
    return round(score, 2)


# ═══════════════════════════════════════════════════════════════════════════
# MAIN PARSING PIPELINE
# ═══════════════════════════════════════════════════════════════════════════

def parse_resume(text: str, extraction_confidence: float = 1.0) -> Dict[str, Any]:
    """
    Complete resume parsing pipeline.
    
    CRITICAL RULES:
    - NEVER invent or estimate data
    - All information must come from the resume text
    - Experience is calculated ONLY from work history dates
    - If data is missing, leave it empty/null
    - Provide confidence scores for quality assessment
    """
    
    # Extract all fields
    name, name_conf = extract_name(text)
    email = extract_email(text)
    phone = extract_phone(text)
    skills = extract_skills(text)
    work_exp, work_conf = extract_work_experience(text)
    education, edu_conf = extract_education(text)
    certifications = extract_certifications(text)
    
    # Calculate experience ONLY from work history
    years_experience = calculate_total_experience_years(work_exp)
    
    # Quality checks
    is_flagged = detect_keyword_stuffing(text, skills)
    flag_reason = "Suspected keyword stuffing or manipulation" if is_flagged else None
    
    # Calculate overall parsing confidence
    parsing_confidence = calculate_parsing_confidence(
        extraction_confidence=extraction_confidence,
        name_confidence=name_conf,
        has_email=bool(email),
        has_phone=bool(phone),
        skills_count=len(skills),
        work_exp_confidence=work_conf,
        edu_confidence=edu_conf,
    )
    
    # Determine if manual review is needed
    needs_manual_review = parsing_confidence < 0.6 or is_flagged or years_experience == 0
    
    # Extract current company and role
    current_company = None
    current_role = None
    if work_exp:
        current_job = work_exp[0]  # Most recent
        if current_job.get("is_current"):
            current_company = current_job.get("company")
            current_role = current_job.get("role")
    
    return {
        "full_name": name,
        "email": email,
        "phone": phone,
        "skills": skills,
        "years_experience": years_experience,
        "education_level": education["level"],
        "education_institution": education["institution"],
        "certifications": certifications,
        "work_experience": work_exp,
        "current_company": current_company,
        "current_role": current_role,
        "is_flagged": is_flagged,
        "flag_reason": flag_reason,
        "parsing_confidence": parsing_confidence,
        "needs_manual_review": needs_manual_review,
        "extraction_quality": {
            "name_confidence": name_conf,
            "has_email": bool(email),
            "has_phone": bool(phone),
            "skills_count": len(skills),
            "work_experience_count": len(work_exp),
            "work_confidence": work_conf,
            "education_confidence": edu_conf,
        }
    }
